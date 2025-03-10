import spacy
import PyPDF2
import docx
import re
import os
from typing import Dict, List, Any, Optional, Union, BinaryIO
from transformers import pipeline
import io
from datetime import datetime

class ResumeParser:
    """
    A class for parsing and extracting information from resumes in PDF or DOCX format.
    """
    
    def __init__(self):
        """Initialize the ResumeParser with NLP models."""
        self.nlp = spacy.load("en_core_web_lg")
        self.ner_pipeline = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english")
        
        # Enhanced skill patterns with categories
        self.skill_patterns = {
            "technical": [
                "python", "java", "javascript", "typescript", "c\\+\\+", "ruby", "php", "swift",
                "kotlin", "rust", "golang", "sql", "nosql", "mongodb", "postgresql", "mysql",
                "oracle", "redis", "elasticsearch", "aws", "azure", "gcp", "docker", "kubernetes",
                "jenkins", "git", "terraform", "ansible", "react", "angular", "vue", "node.js",
                "django", "flask", "spring", "tensorflow", "pytorch", "scikit-learn", "pandas",
                "numpy", "hadoop", "spark", "kafka", "graphql", "rest", "soap", "microservices"
            ],
            "soft": [
                "leadership", "communication", "teamwork", "problem solving", "critical thinking",
                "time management", "project management", "agile", "scrum", "stakeholder management",
                "presentation", "negotiation", "conflict resolution", "mentoring", "coaching"
            ],
            "domain": [
                "machine learning", "artificial intelligence", "data science", "cloud computing",
                "devops", "cybersecurity", "blockchain", "iot", "web development", "mobile development",
                "database administration", "system architecture", "ui/ux design", "business intelligence"
            ]
        }
        
        # Education degree levels and their weights
        self.education_levels = {
            "phd": 1.0,
            "doctorate": 1.0,
            "master": 0.8,
            "mba": 0.8,
            "bachelor": 0.6,
            "associate": 0.4,
            "certification": 0.3
        }
        
        # Experience level patterns
        self.experience_patterns = [
            (r"(\d+)\+?\s*(?:years?|yrs?)", lambda x: int(x)),
            (r"(\d+)-(\d+)\s*(?:years?|yrs?)", lambda x, y: (int(x) + int(y)) / 2),
            (r"(?:since|from)\s*(?:19|20)(\d{2})", lambda x: 2024 - (2000 + int(x)) if int(x) < 24 else 2024 - (1900 + int(x)))
        ]
    
    def extract_text(self, file: Union[str, BinaryIO]) -> str:
        """Extract text from PDF or DOCX file."""
        if isinstance(file, str):
            if file.lower().endswith('.pdf'):
                return self._extract_from_pdf(file)
            elif file.lower().endswith('.docx'):
                return self._extract_from_docx(file)
        else:
            # Handle Streamlit uploaded file
            file_type = file.type
            if file_type == "application/pdf":
                return self._extract_from_pdf_bytes(file)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx_bytes(file)
        raise ValueError("Unsupported file format")

    def _extract_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        return text

    def _extract_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
        return text

    def _extract_from_pdf_bytes(self, file: BinaryIO) -> str:
        """Extract text from PDF bytes."""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        return text

    def _extract_from_docx_bytes(self, file: BinaryIO) -> str:
        """Extract text from DOCX bytes."""
        text = ""
        try:
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
        return text
    
    def extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills with categories from text using advanced NLP."""
        skills = {category: [] for category in self.skill_patterns.keys()}
        doc = self.nlp(text.lower())
        
        # Extract skills using pattern matching and NLP
        for category, patterns in self.skill_patterns.items():
            for skill in patterns:
                if re.search(r'\b' + re.escape(skill) + r'\b', text.lower()):
                    # Verify skill using word embeddings
                    skill_doc = self.nlp(skill)
                    context_window = 10
                    for token in doc:
                        if token.text.lower() == skill.lower():
                            context = doc[max(0, token.i - context_window):min(len(doc), token.i + context_window)]
                            if any(context.similarity(skill_doc) > 0.7 for context in [self.nlp(sent.text) for sent in context.sents]):
                                skills[category].append(skill)
                                break
        
        return {k: list(set(v)) for k, v in skills.items()}
    
    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract detailed education information using NER and pattern matching."""
        education = []
        doc = self.nlp(text)
        
        # Use transformer-based NER for better entity recognition
        ner_results = self.ner_pipeline(text)
        
        # Extract education sections
        education_sections = []
        current_section = ""
        in_education = False
        
        for sent in doc.sents:
            if any(edu_word in sent.text.lower() for edu_word in ["education", "degree", "university", "college"]):
                in_education = True
            elif in_education and len(sent.text.strip()) > 0 and sent.text.strip().isupper():
                in_education = False
            
            if in_education:
                current_section += sent.text + " "
            elif current_section:
                education_sections.append(current_section.strip())
                current_section = ""
        
        if current_section:
            education_sections.append(current_section.strip())
        
        # Process each education section
        for section in education_sections:
            section_doc = self.nlp(section)
            edu_entry = {
                "degree": None,
                "field": None,
                "institution": None,
                "year": None,
                "gpa": None,
                "level": 0.0
            }
            
            # Extract degree and field
            for ent in section_doc.ents:
                if ent.label_ == "DEGREE" or "degree" in ent.text.lower():
                    edu_entry["degree"] = ent.text
                    # Determine education level
                    for level, weight in self.education_levels.items():
                        if level in ent.text.lower():
                            edu_entry["level"] = weight
                            break
                elif ent.label_ == "ORG":
                    edu_entry["institution"] = ent.text
            
            # Extract year
            years = re.findall(r'(19|20)\d{2}', section)
            if years:
                edu_entry["year"] = max(map(int, years))
            
            # Extract GPA
            gpa_match = re.search(r'GPA:?\s*(\d+\.\d+)|(\d+\.\d+)\s*GPA', section)
            if gpa_match:
                edu_entry["gpa"] = float(gpa_match.group(1) or gpa_match.group(2))
            
            if edu_entry["degree"] or edu_entry["institution"]:
                education.append(edu_entry)
        
        return education
    
    def extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract detailed work experience using advanced NLP."""
        experience = []
        doc = self.nlp(text)
        current_year = datetime.now().year
        
        # Extract experience sections
        experience_sections = []
        current_section = ""
        in_experience = False
        
        for sent in doc.sents:
            if any(exp_word in sent.text.lower() for exp_word in ["experience", "employment", "work history"]):
                in_experience = True
            elif in_experience and len(sent.text.strip()) > 0 and sent.text.strip().isupper():
                in_experience = False
            
            if in_experience:
                current_section += sent.text + " "
            elif current_section:
                experience_sections.append(current_section.strip())
                current_section = ""
        
        if current_section:
            experience_sections.append(current_section.strip())
        
        # Process each experience section
        for section in experience_sections:
            section_doc = self.nlp(section)
            exp_entry = {
                "title": None,
                "company": None,
                "duration": None,
                "start_year": None,
                "end_year": None,
                "years": 0.0,
                "responsibilities": [],
                "achievements": []
            }
            
            # Extract job title and company
            for ent in section_doc.ents:
                if ent.label_ == "ORG":
                    exp_entry["company"] = ent.text
                elif ent.label_ in ["TITLE", "PERSON"] and not exp_entry["title"]:
                    exp_entry["title"] = ent.text
            
            # Extract duration and calculate years
            duration_text = section.lower()
            
            # Look for date patterns
            date_patterns = [
                # Present/Current patterns with year
                (r"(20\d{2}|19\d{2})\s*[-–—]\s*(present|current|now|\s*$)", 
                 lambda match: (int(match[0]), current_year)),
                # Year range patterns
                (r"(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2})", 
                 lambda match: (int(match[0]), int(match[1]))),
                # Since year pattern
                (r"since\s*(20\d{2}|19\d{2})", 
                 lambda match: (int(match[0]), current_year))
            ]
            
            # Try to find dates in the text
            for pattern, handler in date_patterns:
                matches = re.findall(pattern, duration_text)
                if matches:
                    try:
                        start_year, end_year = handler(matches[0])
                        exp_entry["start_year"] = start_year
                        exp_entry["end_year"] = end_year
                        exp_entry["duration"] = f"{start_year} - {'Present' if end_year == current_year else end_year}"
                        exp_entry["years"] = end_year - start_year
                        break
                    except Exception as e:
                        print(f"Error processing date pattern: {e}")
                        continue
            
            # If no date pattern found, try experience patterns
            if exp_entry["years"] == 0:
                for pattern, converter in self.experience_patterns:
                    matches = re.findall(pattern, section)
                    if matches:
                        try:
                            if len(matches[0]) == 2:  # Range pattern
                                exp_entry["years"] = converter(matches[0][0], matches[0][1])
                            else:  # Single value pattern
                                exp_entry["years"] = converter(matches[0])
                            break
                        except Exception as e:
                            print(f"Error processing experience pattern: {e}")
                            continue
            
            # Extract responsibilities and achievements
            for sent in section_doc.sents:
                sent_text = sent.text.strip()
                if sent_text.startswith("-") or sent_text.startswith("•"):
                    if any(word in sent_text.lower() for word in ["achieved", "increased", "improved", "reduced", "launched", "created"]):
                        exp_entry["achievements"].append(sent_text.lstrip("- •"))
                    else:
                        exp_entry["responsibilities"].append(sent_text.lstrip("- •"))
            
            if exp_entry["title"] or exp_entry["company"]:
                experience.append(exp_entry)
        
        return experience
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information using regex and NER."""
        contact_info = {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "location": None,
            "website": None
        }
        
        # Extract email
        email_matches = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_matches:
            contact_info["email"] = email_matches[0]
        
        # Extract phone
        phone_matches = re.findall(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        if phone_matches:
            contact_info["phone"] = ''.join(''.join(match) for match in phone_matches[0])
        
        # Extract LinkedIn
        linkedin_matches = re.findall(r'linkedin\.com/in/[A-Za-z0-9_-]+', text)
        if linkedin_matches:
            contact_info["linkedin"] = linkedin_matches[0]
        
        # Extract name using NER
        doc = self.nlp(text[:1000])  # Process first 1000 characters for efficiency
        person_names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if person_names:
            contact_info["name"] = person_names[0]
        
        # Extract location using NER
        locations = [ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]]
        if locations:
            contact_info["location"] = locations[0]
        
        # Extract website
        website_matches = re.findall(r'https?://(?:www\.)?([A-Za-z0-9.-]+\.[A-Za-z]{2,})', text)
        if website_matches:
            contact_info["website"] = website_matches[0]
        
        return contact_info
    
    def calculate_experience_years(self, experience: List[Dict[str, Any]]) -> float:
        """Calculate total years of experience."""
        total_years = 0.0
        for exp in experience:
            total_years += exp.get("years", 0.0)
        return total_years
    
    def parse_resume(self, file: Union[str, BinaryIO]) -> Dict[str, Any]:
        """Parse resume and extract all information."""
        text = self.extract_text(file)
        
        # Extract all information
        contact_info = self.extract_contact_info(text)
        skills = self.extract_skills(text)
        education = self.extract_education(text)
        experience = self.extract_experience(text)
        total_years = self.calculate_experience_years(experience)
        
        # Calculate education level
        max_education_level = max([edu.get("level", 0.0) for edu in education], default=0.0)
        
        return {
            "contact_info": contact_info,
            "skills": skills,
            "education": education,
            "experience": experience,
            "total_years": total_years,
            "education_level": max_education_level,
            "raw_text": text
        } 